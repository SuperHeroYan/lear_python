
import docx
from docx.enum.text import WD_BREAK

doc = docx.Document()

guest_list = open('guests.txt')
read_guest = guest_list.readlines()

for line in range(0, len(read_guest)):
	obj1 = doc.add_paragraph('it would a pleasure to have company')
	# Костбмизация строки
	obj1.style = 'Title'
	obj1.outline = True
	obj1.bold = True
	# Имя гость
	name = read_guest[line]
	obj2 = doc.add_paragraph('Hi %s, youre welcome to party!!'
						% (name) )
	obj2.bold = True
	print(name)
	# Подпись с лбовь
	obj3 = doc.add_paragraph('at 11010 Memory lane on the evning')
	# Костбмизация строки
	obj3.style = 'Title'
	obj3.outline = True
	obj3.bold = True

	obj5 = doc.add_paragraph('at 7 o`clock')
	# Костбмизация строки
	obj5.style = 'Title'
	obj5.outline = True
	obj5.bold = True

	obj5.runs[0].add_break(WD_BREAK.PAGE)

doc.save('twopage.docx')
